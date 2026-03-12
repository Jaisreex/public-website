import os
from bs4 import BeautifulSoup
from docx import Document

# List of target HTML files
files = [
    'index.html',
    'about.html',
    'contact.html',

    'robotics.html',
    'upi-payments.html',

    'laundromat-v2.html'
]

doc = Document()
doc.add_heading('AllZiva Website Content Compilation', 0)

for file in files:
    if os.path.exists(file):
        with open(file, 'r', encoding='utf-8') as f:
            html_content = f.read()

        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Add a heading for the page
        doc.add_heading(f'Page: {file.replace(".html", "").replace("-", " ").title()}', level=1)

        # Basic parsing: iterate through body tags and extract text
        if soup.body:
            for element in soup.body.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'li']):
                text = element.get_text(strip=True)
                if not text: continue
                # Do not get nested text twice
                # We can just check if parent is also in the list, but it's simpler to just output the element.
                # However find_all gets all nested elements too.
                # Let's limit finding to elements that usually don't nest the same type.
                if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                    doc.add_heading(text, level=int(element.name[1]))
                elif element.name == 'p':
                    doc.add_paragraph(text)
                elif element.name == 'li':
                    doc.add_paragraph(text, style='List Bullet')
        else:
             doc.add_paragraph("No body content found.")
            
        doc.add_page_break()

doc.save('AllZiva_Website_Content.docx')
print("Successfully created AllZiva_Website_Content.docx")
